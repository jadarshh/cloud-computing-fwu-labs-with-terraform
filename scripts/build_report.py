"""Build the Cloud Computing Lab Report (.docx) for the student named in config.yaml.

Reads `config.yaml` at the repo root and screenshots from `screenshots/lab-NN/`,
writes a single Word document to `report/<Full_Name>_Cloud_Computing_Lab_Report.docx`.

Run:
    python scripts/build_report.py
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from lib.config import load  # noqa: E402

ROOT = Path(__file__).resolve().parent.parent
SHOTS = ROOT / "screenshots"

# ---------------------------------------------------------------------------
# Visual design tokens
# ---------------------------------------------------------------------------

COLOR_PRIMARY = RGBColor(0x1F, 0x38, 0x64)   # deep academic blue (Heading 1)
COLOR_SECONDARY = RGBColor(0x2E, 0x74, 0xB5) # mid blue (Heading 3, accents)
COLOR_GREY = RGBColor(0x59, 0x59, 0x59)      # captions, sub-text
COLOR_RULE = RGBColor(0xBF, 0xBF, 0xBF)      # horizontal rule under lab titles

FONT_BODY = "Calibri"
FONT_HEADING = "Calibri"

SIZE_BODY = 11
SIZE_LAB_TITLE = 22
SIZE_H1 = 14
SIZE_H3 = 12
SIZE_CAPTION = 10
SIZE_TOC_ENTRY = 12
SIZE_COVER_TITLE = 20   # university line — was 30
SIZE_COVER_BIG = 22     # subject (CLOUD COMPUTING) — unchanged
SIZE_COVER_NAME = 15    # student full name — was 18


def output_path(cfg) -> Path:
    fname = cfg.student.full_name.replace(" ", "_") + "_Cloud_Computing_Lab_Report.docx"
    return ROOT / "report" / fname


# ---------------------------------------------------------------------------
# Low-level helpers (XML / styling)
# ---------------------------------------------------------------------------

def _set_run(run, *, bold=False, italic=False, size=None, color=None, font_name=None):
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if font_name is not None:
        run.font.name = font_name


def _para_spacing(p, *, space_before=0, space_after=4, line=1.15, keep_with_next=False, keep_together=False, page_break_before=False):
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.widow_control = True
    if keep_with_next:
        pf.keep_with_next = True
    if keep_together:
        pf.keep_together = True
    if page_break_before:
        pf.page_break_before = True


def _add_horizontal_rule(paragraph, color=COLOR_RULE, size=8):
    """Add a bottom border to a paragraph as a horizontal rule under it."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(*color))
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_page_field(run):
    """Insert {PAGE} field into a run for live page numbers."""
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _add_pages_field(run):
    """Insert {NUMPAGES} field into a run for live total page count."""
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "NUMPAGES"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_heading_styles(doc):
    """Configure Word's native Heading 1 / Heading 2 styles to match the
    polished reference (navy bold, sized for an academic doc).

    Using real heading styles (instead of raw runs) is what lets Word's
    navigation pane jump between labs and lets the TOC field populate
    page numbers automatically.
    """
    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT_HEADING
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = COLOR_PRIMARY
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT_HEADING
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = COLOR_PRIMARY
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(2)
    h2.paragraph_format.keep_with_next = True


# ---------------------------------------------------------------------------
# Section / page setup
# ---------------------------------------------------------------------------

def configure_sections(doc, cfg):
    """Set Moderate margins and configure header/footer with page numbers."""
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    section.different_first_page_header_footer = True

    header = section.header
    h_para = header.paragraphs[0]
    h_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = h_para.add_run(
        f"{cfg.subject.name} Lab Report  |  {cfg.student.full_name}  |  Roll No. {cfg.student.roll_number}"
    )
    _set_run(hr, italic=True, size=10, color=COLOR_GREY)
    _add_horizontal_rule(h_para, color=COLOR_RULE, size=4)

    footer = section.footer
    f_para = footer.paragraphs[0]
    f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    label = f_para.add_run("Page ")
    _set_run(label, size=10, color=COLOR_GREY)

    page_run = f_para.add_run()
    _set_run(page_run, size=10, color=COLOR_GREY)
    _add_page_field(page_run)

    of_run = f_para.add_run(" of ")
    _set_run(of_run, size=10, color=COLOR_GREY)

    pages_run = f_para.add_run()
    _set_run(pages_run, size=10, color=COLOR_GREY)
    _add_pages_field(pages_run)


def configure_default_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(SIZE_BODY)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.3


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def _cover_centered(doc, text, *, bold=False, italic=False, size=None, color=None,
                    space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p, space_before=space_before, space_after=space_after, line=1.2)
    r = p.add_run(text)
    _set_run(r, bold=bold, italic=italic, size=size, color=color)
    return p


def add_cover_page(doc, cfg):
    for _ in range(2):
        doc.add_paragraph()

    _cover_centered(doc, cfg.institution.university.upper(),
                    bold=True, size=SIZE_COVER_TITLE, color=COLOR_PRIMARY,
                    space_after=8)
    _cover_centered(doc, cfg.institution.faculty,
                    bold=False, size=14, color=COLOR_PRIMARY)
    p_school = _cover_centered(doc, cfg.institution.school,
                               italic=True, size=13, color=COLOR_GREY)
    _cover_centered(doc, cfg.institution.location,
                    italic=True, size=11, color=COLOR_GREY)
    _add_horizontal_rule(p_school, color=COLOR_PRIMARY, size=12)

    for _ in range(3):
        doc.add_paragraph()

    _cover_centered(doc, "A LAB REPORT",
                    bold=True, size=14, color=COLOR_GREY, space_after=4)
    _cover_centered(doc, "ON",
                    size=12, color=COLOR_GREY, space_after=4)
    p_title = _cover_centered(doc, cfg.subject.name.upper(),
                              bold=True, size=SIZE_COVER_BIG,
                              color=COLOR_PRIMARY, space_after=2)
    _cover_centered(doc, "(Practical)",
                    italic=True, size=12, color=COLOR_GREY, space_after=4)
    _add_horizontal_rule(p_title, color=COLOR_PRIMARY, size=8)

    for _ in range(3):
        doc.add_paragraph()

    _cover_centered(doc, "Submitted by:",
                    italic=True, size=12, color=COLOR_GREY, space_after=4)
    _cover_centered(doc, cfg.student.full_name,
                    bold=True, size=SIZE_COVER_NAME,
                    color=COLOR_PRIMARY, space_after=2)
    _cover_centered(doc, f"Roll No. {cfg.student.roll_number}",
                    size=12, color=COLOR_GREY, space_after=2)
    _cover_centered(doc, f"{cfg.student.semester} Semester  |  {cfg.student.program}",
                    size=11, color=COLOR_GREY, space_after=4)

    for _ in range(2):
        doc.add_paragraph()

    _cover_centered(doc, "Submitted to:",
                    italic=True, size=12, color=COLOR_GREY, space_after=4)
    _cover_centered(doc, cfg.professor.name,
                    bold=True, size=13, color=COLOR_PRIMARY, space_after=2)

    # Split professor.title on the LAST comma so the trailing institution name
    # becomes its own line. If no comma, the whole title goes on one line.
    title = cfg.professor.title
    if ", " in title:
        title_line_1, title_line_2 = title.rsplit(", ", 1)
    else:
        title_line_1, title_line_2 = title, ""
    _cover_centered(doc, title_line_1,
                    size=11, color=COLOR_GREY, space_after=2)
    if title_line_2:
        _cover_centered(doc, title_line_2,
                        size=11, color=COLOR_GREY, space_after=4)

    for _ in range(3):
        doc.add_paragraph()

    _cover_centered(doc, "Signature:  ____________________",
                    size=11, color=COLOR_GREY, space_after=4)
    _cover_centered(doc, "Date:  ____________________",
                    size=11, color=COLOR_GREY, space_after=2)

    doc.add_page_break()


def add_acknowledgement(doc, cfg):
    """Acknowledgement page: Heading 1 'ACKNOWLEDGEMENT', three justified body
    paragraphs, signed at the bottom right with name, roll, and program.
    """
    s = cfg.student
    inst = cfg.institution
    prof = cfg.professor

    h = doc.add_paragraph(style="Heading 1")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.add_run("ACKNOWLEDGEMENT")

    paragraphs = [
        (
            f"I would like to express my sincere gratitude to my subject teacher, "
            f"{prof.name}, for guiding me throughout the {cfg.subject.name} "
            f"laboratory work and for the patient explanations of cloud concepts "
            f"that helped me understand the material well beyond what a textbook "
            f"alone could offer."
        ),
        (
            f"I am also thankful to the {inst.faculty} and the {inst.school} at "
            f"{inst.university} for providing the lab environment and the "
            f"opportunity to complete the practical exercises that make up this "
            f"report."
        ),
        (
            f"Finally, I would like to thank my classmates of {s.semester} semester "
            f"{s.program} for the discussions and the willingness to compare notes "
            f"while working through the AWS console — those small conversations "
            f"often saved hours of confusion."
        ),
    ]
    for txt in paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _para_spacing(p, space_before=0, space_after=8, line=1.3)
        r = p.add_run(txt)
        _set_run(r, size=SIZE_BODY)

    # Signed block, right-aligned
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _para_spacing(p, space_before=14, space_after=2, line=1.2)
    r = p.add_run(s.full_name)
    _set_run(r, bold=True, size=12, color=COLOR_PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _para_spacing(p, space_before=0, space_after=2, line=1.2)
    r = p.add_run(f"Roll No. {s.roll_number}")
    _set_run(r, size=11, color=COLOR_GREY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _para_spacing(p, space_before=0, space_after=2, line=1.2)
    r = p.add_run(f"{s.program}, {s.semester} Semester")
    _set_run(r, size=11, color=COLOR_GREY)

    doc.add_page_break()


def _add_toc_field(paragraph, instr=r'TOC \o "1-1" \h \z \u'):
    """Emit a Word TOC field that builds a ToC from Heading 1 entries."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    run._r.append(fld_begin)

    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    run._r.append(instr_el)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    placeholder = paragraph.add_run(
        "Right-click and select \"Update Field\" to populate."
    )
    _set_run(placeholder, italic=True, size=10, color=COLOR_GREY)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    placeholder._r.append(fld_end)


def _enable_update_fields_on_open(doc):
    """Set w:updateFields in settings.xml so Word refreshes fields on open."""
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)
    else:
        existing.set(qn("w:val"), "true")


# ---------------------------------------------------------------------------
# Table of contents
# ---------------------------------------------------------------------------

def add_toc(doc, labs):
    """Add a centered TABLE OF CONTENTS heading and a Word TOC field.

    The `labs` argument is unused (field reads from document heading structure)
    but kept for call-site compatibility.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p, space_before=4, space_after=8, keep_with_next=True)
    r = p.add_run("TABLE OF CONTENTS")
    _set_run(r, bold=True, size=18, color=COLOR_PRIMARY)
    _add_horizontal_rule(p, color=COLOR_PRIMARY, size=8)

    toc_para = doc.add_paragraph()
    _para_spacing(toc_para, space_before=4, space_after=8, line=1.2)
    _add_toc_field(toc_para)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Per-lab elements
# ---------------------------------------------------------------------------

def _add_h1(doc, text, *, keep_with_next=True):
    """Per-lab section heading using Word Heading 2 style.

    Named _add_h1 for historical reasons; visually these are H2 because
    the lab title itself is H1.
    """
    p = doc.add_paragraph(style="Heading 2")
    _para_spacing(p, space_before=10, space_after=4, line=1.2,
                  keep_with_next=keep_with_next)
    p.add_run(text)
    return p


def _add_h3(doc, text, *, keep_with_next=True):
    """Step sub-heading inside a lab procedure. Plain black bold 12pt."""
    p = doc.add_paragraph()
    _para_spacing(p, space_before=6, space_after=2, line=1.2,
                  keep_with_next=keep_with_next)
    r = p.add_run(text)
    _set_run(r, bold=True, size=SIZE_H3)
    return p


def _add_body(doc, text, *, justify=True, keep_together=False):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _para_spacing(p, space_before=0, space_after=6, line=1.3,
                  keep_together=keep_together)
    r = p.add_run(text)
    _set_run(r, size=SIZE_BODY)
    return p


def _add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    _para_spacing(p, space_before=0, space_after=2, line=1.2)
    if not p.text.strip():
        r = p.add_run(text)
        _set_run(r, size=SIZE_BODY)
    else:
        for r in p.runs:
            r.text = ""
        r = p.add_run(text)
        _set_run(r, size=SIZE_BODY)


def _set_cell_borders(cell, color="BFBFBF", size="4"):
    """Add a single 1pt light-grey border around a table cell on all four sides."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        borders.append(b)


def _add_services_table(doc, services):
    """Render services as a clean 2-column bordered grid (no bullets)."""
    n = len(services)
    rows = (n + 1) // 2
    table = doc.add_table(rows=rows, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, svc in enumerate(services):
        r, c = divmod(idx, 2)
        cell = table.cell(r, c)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        para = cell.paragraphs[0]
        _para_spacing(para, space_before=2, space_after=2, line=1.15)
        run = para.add_run(svc)
        _set_run(run, size=SIZE_BODY)
        _set_cell_borders(cell)

    # Give the trailing empty cell a border too if services count is odd
    if n % 2 == 1:
        empty_cell = table.cell(rows - 1, 1)
        _set_cell_borders(empty_cell)

    after = doc.add_paragraph()
    _para_spacing(after, space_before=0, space_after=4, line=1.0)


def _add_screenshot(doc, image_path, caption):
    if not Path(image_path).exists():
        warn = doc.add_paragraph()
        wr = warn.add_run(f"[MISSING SCREENSHOT: {image_path}]")
        _set_run(wr, italic=True, color=RGBColor(0xCC, 0x00, 0x00))
        return

    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(pic_para, space_before=4, space_after=2, line=1.0,
                  keep_with_next=True, keep_together=True)
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Inches(5.6))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(cap, space_before=0, space_after=8, line=1.15,
                  keep_together=True)
    cr = cap.add_run(caption)
    _set_run(cr, bold=True, italic=True, size=SIZE_CAPTION, color=COLOR_GREY)


def add_lab(doc, n, lab, *, is_first=False):
    title_p = doc.add_paragraph(style="Heading 1")
    _para_spacing(title_p, space_before=0, space_after=8, line=1.2,
                  keep_with_next=True,
                  page_break_before=not is_first)
    title_p.add_run(f"Lab {n}: {lab['title']}")
    _add_horizontal_rule(title_p, color=COLOR_PRIMARY, size=12)

    _add_h1(doc, "Objective")
    _add_body(doc, lab["objective"])

    _add_h1(doc, "AWS Services Used")
    _add_services_table(doc, lab["services"])

    _add_h1(doc, "Step-by-Step Procedure")
    for step_title, lines in lab["procedure"]:
        _add_h3(doc, step_title)
        for line in lines:
            _add_bullet(doc, line)

    _add_h1(doc, "Screenshots")
    for path, caption in lab["screenshots"]:
        _add_screenshot(doc, SHOTS / path, caption)

    _add_h1(doc, "Observations and Results")
    _add_body(doc, lab["observations"])


def add_conclusion(doc, cfg):
    """Closing Conclusion page. Heading 1 'Conclusion' (joins the ToC) +
    five justified paragraphs synthesising the work across the nine labs.
    """
    s = cfg.student

    h = doc.add_paragraph(style="Heading 1")
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h.add_run("Conclusion")
    _para_spacing(h, space_before=0, space_after=8, line=1.2,
                  page_break_before=True)

    paragraphs = [
        (
            "Across the nine labs in this report, I worked through the practical "
            "building blocks of a typical AWS deployment, starting from the "
            "network layer at the bottom and finishing with managed services and "
            "infrastructure as code at the top. The labs were small enough to "
            "finish in a single session, but together they cover most of what a "
            "real cloud workload would touch."
        ),
        (
            "Labs 1 and 4 covered the network foundation — a VPC with public and "
            "private subnets, an internet gateway, route tables, and tiered "
            "security groups. The most useful idea here was that everything else "
            "in AWS sits on top of this network, so getting the routing and "
            "security boundaries right makes every later step easier."
        ),
        (
            "Labs 2, 3, and 5 covered compute and traffic delivery — an EC2 "
            "instance with a startup script, S3 hosting a static website without "
            "any server, and an Application Load Balancer in front of an Auto "
            "Scaling Group across two availability zones. Watching the ALB "
            "alternate between backend instances on each refresh made the "
            "horizontal-scaling story click in a way that reading about it never "
            "did."
        ),
        (
            "Lab 6 covered identity and access — IAM users, groups, and managed "
            "policies — and made the difference between assigning permissions "
            "through a group and attaching them directly to a user concrete. "
            "Group-based permissioning is what scales in practice, and the "
            "console makes it obvious why."
        ),
        (
            f"Labs 7, 8, and 9 covered the higher-level managed services and "
            f"infrastructure as code: a Lambda function writing into a DynamoDB "
            f"table on each invocation, a Fargate task running an Nginx container "
            f"without any EC2 host to manage, an SNS topic fanning out a single "
            f"publish to both an SQS queue and an email subscriber, and finally a "
            f"CloudFormation stack creating a VPC and an S3 bucket from a single "
            f"YAML file. Beyond the AWS services themselves, the labs gave me "
            f"practical experience with the AWS Management Console, the AWS CLI, "
            f"basic shell user-data scripts, and the discipline of capturing "
            f"evidence at each step — habits that will carry over into any cloud "
            f"work I do after this {s.semester} semester."
        ),
    ]
    for txt in paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _para_spacing(p, space_before=0, space_after=8, line=1.3)
        r = p.add_run(txt)
        _set_run(r, size=SIZE_BODY)


# ---------------------------------------------------------------------------
# Lab content (functions of cfg)
# ---------------------------------------------------------------------------

def _region_short(label: str) -> str:
    """Extract the short location part from 'Asia Pacific (Mumbai)' -> 'Mumbai'."""
    if "(" in label and ")" in label:
        return label[label.index("(") + 1:label.index(")")]
    return label


def lab_1(cfg):
    s = cfg.student
    aws = cfg.aws
    region_short = _region_short(aws.region_label)
    return {
        "title": "Virtual Cloud Environment (VPC)",
        "objective": (
            f"Create and configure a Virtual Private Cloud (VPC) on AWS in the "
            f"{region_short} region. The VPC acts as the base private network within "
            "which other lab resources can later be launched in isolation from the "
            "rest of the internet."
        ),
        "services": ["Amazon VPC", "AWS Management Console"],
        "procedure": [
            ("Step 1: Sign in and select region", [
                "Open https://console.aws.amazon.com and sign in with the IAM user.",
                f"From the top-right region dropdown, select {aws.region_label} - {aws.region}.",
            ]),
            ("Step 2: Open the VPC service", [
                "Search for 'VPC' in the search bar and open the VPC dashboard.",
                "From the left navigation, click on Your VPCs.",
                "Click on the Create VPC button.",
            ]),
            ("Step 3: Configure the VPC", [
                "Select 'VPC only' as the resource to create.",
                f"Set Name tag to vpc-{s.suffix}.",
                "Set IPv4 CIDR block to 10.20.0.0/16.",
                "Leave IPv6 CIDR block as 'No IPv6 CIDR block' and Tenancy as Default.",
                "Click Create VPC.",
            ]),
            ("Step 4: Verify creation", [
                "Confirm the success banner appears with the new VPC ID.",
                "Verify that the State is Available and DNS resolution is Enabled.",
                "Note down the VPC ID for reference.",
            ]),
        ],
        "screenshots": [
            ("lab-01/1.1-vpc-list.png",
             f"Figure 1.1 — VPC vpc-{s.suffix} listed in the {region_short} region with State = Available."),
            ("lab-01/1.2-vpc-details.png",
             "Figure 1.2 — Details panel showing CIDR 10.20.0.0/16 and DNS settings enabled."),
        ],
        "observations": (
            f"The VPC vpc-{s.suffix} was created successfully in the {aws.region} region with "
            "the IPv4 CIDR block 10.20.0.0/16. DNS resolution and DNS hostnames were "
            "enabled by default. The VPC entered the Available state immediately after "
            "creation, confirming that the base network for the remaining labs has been "
            "set up correctly."
        ),
    }


def lab_2(cfg):
    s = cfg.student
    return {
        "title": "Compute Instances and Startup Scripts (EC2)",
        "objective": (
            "Launch an Amazon EC2 instance and use a startup script (user_data) to "
            "automatically install the Apache web server and serve a custom welcome "
            "page on port 80 the moment the instance boots."
        ),
        "services": [
            "Amazon EC2",
            "Amazon Machine Image (Amazon Linux 2023)",
            "Security Groups",
            "AWS Management Console",
        ],
        "procedure": [
            ("Step 1: Sign in and open EC2 console", [
                f"Sign in to the AWS Console and select region {cfg.aws.region_label} - {cfg.aws.region}.",
                "Open the EC2 service from the search bar and click Launch instance.",
            ]),
            ("Step 2: Configure the instance", [
                f"Set Name to ec2-{s.suffix}.",
                "Choose AMI: Amazon Linux 2023 (64-bit x86, free tier eligible).",
                "Select Instance type: t2.micro (free tier eligible).",
                "Skip the key pair section (or proceed without one for this lab).",
            ]),
            ("Step 3: Configure security group", [
                f"Create a new security group named {s.suffix}-ec2-sg.",
                "Add inbound rule: SSH on port 22, source 0.0.0.0/0.",
                "Add inbound rule: HTTP on port 80, source 0.0.0.0/0.",
                "Leave outbound as default (all traffic allowed).",
            ]),
            ("Step 4: Provide user data startup script", [
                "Expand the Advanced details section.",
                "In the User data field paste the bash script that installs httpd and "
                f"writes the custom index.html (welcome page mentioning {s.display_name}, Roll No. {s.roll_number}).",
                "Click Launch instance.",
            ]),
            ("Step 5: Verify and test", [
                "After about 60 to 90 seconds, refresh the Instances page until the "
                "instance shows State = Running and 2/2 status checks passed.",
                "Copy the Public IPv4 address.",
                "Open http://<public-ip> in a browser to confirm Apache is serving "
                "the welcome page.",
            ]),
        ],
        "screenshots": [
            ("lab-02/2.1-ec2-running.png",
             f"Figure 2.1 — EC2 instance ec2-{s.suffix} in the Running state with status checks passed."),
            ("lab-02/2.2-ec2-details.png",
             f"Figure 2.2 — Instance details: t2.micro, Amazon Linux 2023, attached security group {s.suffix}-ec2-sg."),
            ("lab-02/2.3-browser-welcome.png",
             "Figure 2.3 — Browser at the public IP showing the custom welcome page served by Apache."),
        ],
        "observations": (
            f"The EC2 instance ec2-{s.suffix} booted successfully in the {cfg.aws.region} region. "
            "The user_data script ran during the first boot and installed Apache (httpd), "
            "started the service, and wrote a custom index.html. Once the instance "
            "reached the Running state, the public IP responded over HTTP with the "
            "welcome page, which confirmed that the startup script executed correctly "
            "and the security group allowed inbound HTTP traffic."
        ),
    }


def lab_3(cfg):
    s = cfg.student
    region_short = _region_short(cfg.aws.region_label)
    return {
        "title": "Object Storage and Static Website Hosting (S3)",
        "objective": (
            "Create an Amazon S3 bucket and configure it for static website hosting. "
            "Upload an HTML page and access the rendered site through the S3 website "
            "endpoint without using any web server."
        ),
        "services": [
            "Amazon S3",
            "S3 Static Website Hosting",
            "S3 Bucket Policy",
            "AWS Management Console",
        ],
        "procedure": [
            ("Step 1: Create the bucket", [
                "Open the S3 service in the AWS Console.",
                "Click Create bucket.",
                f"Enter a globally unique bucket name ({s.suffix}-static-<random>) and "
                f"choose AWS Region: {cfg.aws.region_label} {cfg.aws.region}.",
            ]),
            ("Step 2: Allow public access", [
                "In Block Public Access settings, uncheck 'Block all public access'.",
                "Acknowledge the warning by ticking the checkbox.",
                "Click Create bucket.",
            ]),
            ("Step 3: Enable static website hosting", [
                "Open the bucket and go to the Properties tab.",
                "Scroll to Static website hosting and click Edit.",
                "Choose 'Enable', set Hosting type to 'Host a static website'.",
                "Set Index document to index.html and Error document to error.html.",
                "Save changes.",
            ]),
            ("Step 4: Set bucket policy and upload files", [
                "Go to the Permissions tab and add a bucket policy that allows "
                "s3:GetObject for everyone (Principal *).",
                "Go to the Objects tab and upload index.html and error.html.",
            ]),
            ("Step 5: Access the website", [
                "Return to Properties → Static website hosting and copy the bucket "
                "website endpoint URL.",
                "Open the endpoint in a browser to view the rendered page.",
            ]),
        ],
        "screenshots": [
            ("lab-03/3.1-bucket-overview.png",
             f"Figure 3.1 — S3 bucket created in the {region_short} region with the unique {s.suffix}-static name."),
            ("lab-03/3.2-static-hosting.png",
             "Figure 3.2 — Static website hosting enabled in bucket properties with the website endpoint."),
            ("lab-03/3.3-browser-site.png",
             "Figure 3.3 — The static website rendered in a browser via the S3 website endpoint."),
        ],
        "observations": (
            "The S3 bucket was created with a globally unique name and configured to "
            "host a static website. After disabling Block Public Access and adding "
            "a bucket policy that allows public read on objects, the index.html and "
            "error.html files were uploaded. The bucket website endpoint served the "
            "page over HTTP and rendered correctly in the browser, demonstrating "
            "that S3 alone can host a fully static site without any backend server."
        ),
    }


def lab_4(cfg):
    s = cfg.student
    aws = cfg.aws
    return {
        "title": "Virtual Networking — Subnets, Routing, and Security Groups",
        "objective": (
            "Build a custom VPC with one public subnet and one private subnet across "
            "two availability zones, attach an internet gateway, configure route "
            "tables, and define two tiered security groups (web and database) to "
            "demonstrate a typical multi-tier network design."
        ),
        "services": [
            "Amazon VPC",
            "Subnets (Public and Private)",
            "Internet Gateway",
            "Route Tables",
            "Security Groups",
        ],
        "procedure": [
            ("Step 1: Create the custom VPC", [
                "Open the VPC dashboard.",
                f"Create a new VPC named vpc-{s.suffix}-net with CIDR 10.30.0.0/16.",
                "Enable DNS hostnames and DNS resolution.",
            ]),
            ("Step 2: Create subnets", [
                f"Create a public subnet subnet-{s.suffix}-public with CIDR 10.30.1.0/24 "
                f"in availability zone {aws.region}a.",
                "Enable Auto-assign public IPv4 address on the public subnet.",
                f"Create a private subnet subnet-{s.suffix}-private with CIDR 10.30.2.0/24 "
                f"in availability zone {aws.region}b.",
            ]),
            ("Step 3: Attach an internet gateway", [
                f"Create an Internet Gateway named igw-{s.suffix}.",
                f"Attach the IGW to vpc-{s.suffix}-net.",
            ]),
            ("Step 4: Configure route tables", [
                f"Create a public route table rt-{s.suffix}-public in vpc-{s.suffix}-net.",
                f"Add a route 0.0.0.0/0 -> igw-{s.suffix}.",
                f"Associate rt-{s.suffix}-public with subnet-{s.suffix}-public.",
                f"Create a private route table rt-{s.suffix}-private and associate it with "
                f"subnet-{s.suffix}-private (no internet route).",
            ]),
            ("Step 5: Create security groups", [
                f"Create {s.suffix}-web-sg in vpc-{s.suffix}-net allowing inbound HTTP (80) and "
                "HTTPS (443) from 0.0.0.0/0.",
                f"Create {s.suffix}-db-sg in vpc-{s.suffix}-net allowing inbound MySQL (3306) "
                f"only from the source {s.suffix}-web-sg, not from the open internet.",
            ]),
        ],
        "screenshots": [
            ("lab-04/4.1-vpc-overview.png",
             f"Figure 4.1 — Custom VPC vpc-{s.suffix}-net details with CIDR 10.30.0.0/16."),
            ("lab-04/4.2-subnets.png",
             "Figure 4.2 — Public and private subnets in two availability zones."),
            ("lab-04/4.3-route-table-public.png",
             "Figure 4.3 — Public route table with the 0.0.0.0/0 route pointing to the internet gateway."),
            ("lab-04/4.4-igw-attached.png",
             f"Figure 4.4 — Internet gateway igw-{s.suffix} attached to vpc-{s.suffix}-net."),
            ("lab-04/4.5-sg-web-rules.png",
             f"Figure 4.5 — {s.suffix}-web-sg inbound rules allowing HTTP and HTTPS from anywhere."),
            ("lab-04/4.6-sg-db-rules.png",
             f"Figure 4.6 — {s.suffix}-db-sg allowing MySQL only from the {s.suffix}-web-sg source."),
        ],
        "observations": (
            "A custom multi-tier VPC was successfully built. The public subnet is "
            "internet-facing through the IGW and its associated public route table, "
            "while the private subnet has no direct internet route, which is the "
            "expected behaviour for backend or database tiers. The two security "
            "groups model a typical web-to-database trust boundary - the web tier "
            "is reachable from the public internet, but the database tier accepts "
            "MySQL traffic only from the web SG. This setup forms the standard "
            "starting point for any production-style AWS workload."
        ),
    }


def lab_5(cfg):
    s = cfg.student
    aws = cfg.aws
    return {
        "title": "Load Balancer and Auto-Scaling Simulation",
        "objective": (
            "Configure an Application Load Balancer (ALB) in front of an Auto Scaling "
            "Group (ASG) of EC2 instances spread across two availability zones, "
            "demonstrating horizontal scaling and even distribution of incoming HTTP "
            "traffic."
        ),
        "services": [
            "Application Load Balancer (ALB)",
            "Target Groups",
            "Launch Templates",
            "Auto Scaling Groups",
            "Amazon EC2",
        ],
        "procedure": [
            ("Step 1: Build the supporting VPC", [
                f"Create a VPC vpc-{s.suffix}-alb with CIDR 10.40.0.0/16.",
                f"Create two public subnets in {aws.region}a and {aws.region}b with auto "
                "public IP enabled.",
                "Attach an internet gateway and add the default route to the route table.",
            ]),
            ("Step 2: Create launch template", [
                f"Create launch template lt-{s.suffix} using Amazon Linux 2023 and t2.micro.",
                "Provide a user-data script that installs Apache and writes an index.html "
                "showing the EC2 instance ID and availability zone.",
                f"Attach the security group {s.suffix}-alb-web-sg (inbound HTTP from anywhere).",
            ]),
            ("Step 3: Create target group and ALB", [
                f"Create target group tg-{s.suffix}: HTTP, port 80, target type Instance, "
                "with health check path '/'.",
                f"Create an internet-facing Application Load Balancer alb-{s.suffix} in "
                "the two public subnets.",
                f"Add an HTTP:80 listener that forwards to tg-{s.suffix}.",
            ]),
            ("Step 4: Create the Auto Scaling Group", [
                f"Create asg-{s.suffix} referencing the launch template above.",
                "Set min size = 2, desired = 2, max size = 3.",
                f"Attach to the target group tg-{s.suffix} and use VPC zone identifier "
                "with both public subnets.",
                "Set health check type to ELB.",
            ]),
            ("Step 5: Verify load balancing", [
                "Wait for the two ASG instances to pass target group health checks.",
                "Open the ALB DNS name in a browser.",
                "Refresh the page multiple times - each refresh should sometimes "
                "return a different instance ID, proving that the ALB distributes "
                "requests across both backend instances.",
            ]),
        ],
        "screenshots": [
            ("lab-05/5.1-alb-overview.png",
             f"Figure 5.1 — Application Load Balancer alb-{s.suffix} with internet-facing scheme and DNS name."),
            ("lab-05/5.2-asg-instances.png",
             f"Figure 5.2 — Auto Scaling Group asg-{s.suffix} with two InService instances."),
            ("lab-05/5.3-target-group-health.png",
             f"Figure 5.3 — Target group tg-{s.suffix} showing both targets healthy."),
            ("lab-05/5.4-browser-alb-1.png",
             "Figure 5.4 — First request via the ALB DNS, served by the first ASG instance."),
            ("lab-05/5.5-browser-alb-2.png",
             "Figure 5.5 — After a refresh, the request is served by the second ASG instance, confirming load balancing."),
        ],
        "observations": (
            "The ALB and ASG combination worked as expected. Both instances launched "
            "by the ASG passed the target-group health check within roughly two "
            "minutes of boot, after which the ALB started routing traffic to them. "
            "Hitting the ALB DNS name and refreshing the page returned different "
            "instance IDs across requests, which confirmed that the load balancer "
            "was distributing traffic across both backend hosts. The ASG also "
            "demonstrates self-healing - if either instance is terminated, the ASG "
            "would launch a replacement to maintain the desired capacity of two."
        ),
    }


def lab_6(cfg):
    s = cfg.student
    return {
        "title": "IAM Users, Groups and Policy Configuration",
        "objective": (
            "Create IAM users, an IAM group, and attach AWS managed policies to "
            "demonstrate role-based access control - one developer user with EC2 "
            "permissions through group membership, and one read-only user with a "
            "policy attached directly."
        ),
        "services": [
            "AWS Identity and Access Management (IAM)",
            "IAM Users",
            "IAM Groups",
            "AWS Managed Policies",
        ],
        "procedure": [
            ("Step 1: Create the IAM group", [
                "Open the IAM console.",
                "Go to User groups and click Create group.",
                f"Set group name to Developers-{s.suffix}.",
                "Attach the AWS managed policy AmazonEC2FullAccess to the group.",
            ]),
            ("Step 2: Create the developer user", [
                "Go to Users and click Create user.",
                f"Set user name to {s.suffix}-dev.",
                "Enable AWS Management Console access and let the system generate a password.",
                "Untick 'Require password reset on next sign-in' for this lab.",
                f"Add the user to the Developers-{s.suffix} group on the next screen.",
                "Complete the wizard and note down the auto-generated password.",
            ]),
            ("Step 3: Create the read-only user", [
                f"Create another user named {s.suffix}-readonly.",
                "On the Set permissions screen choose 'Attach policies directly'.",
                "Search for and attach the AWS managed policy ReadOnlyAccess.",
                "Complete the wizard.",
            ]),
            ("Step 4: Sign in as the IAM user", [
                "Open an Incognito/private browser window.",
                "Go to the account-specific console URL: "
                "https://<account-id>.signin.aws.amazon.com/console.",
                f"Sign in with the IAM user name {s.suffix}-dev and its password.",
                f"Verify the top-right banner shows {s.suffix}-dev signed in.",
            ]),
        ],
        "screenshots": [
            ("lab-06/6.1-iam-users-list.png",
             f"Figure 6.1 — IAM Users page showing both {s.suffix}-dev and {s.suffix}-readonly."),
            ("lab-06/6.2-iam-group-developers.png",
             f"Figure 6.2 — The Developers-{s.suffix} group with AmazonEC2FullAccess attached."),
            ("lab-06/6.3-iam-user-readonly-policies.png",
             f"Figure 6.3 — {s.suffix}-readonly user with the ReadOnlyAccess policy attached directly."),
            ("lab-06/6.4-iam-login-as-dev.png",
             f"Figure 6.4 — Successful sign-in as {s.suffix}-dev confirmed by the top-right banner."),
        ],
        "observations": (
            f"Two IAM users with different access patterns were created. {s.suffix}-dev "
            "received its EC2 permissions indirectly, by being a member of the "
            f"Developers-{s.suffix} group which had AmazonEC2FullAccess attached. "
            f"{s.suffix}-readonly received its permissions directly through a policy "
            "attached at the user level. The sign-in test in an incognito window "
            "succeeded with the auto-generated password, confirming that the login "
            "profile was created correctly. Both styles - group-based and "
            "user-attached - are valid, but in real environments the group-based "
            "pattern scales much better because permission changes apply to all "
            "members at once."
        ),
    }


def lab_7(cfg):
    s = cfg.student
    return {
        "title": "Serverless Function Deployment (Lambda, Fargate, DynamoDB)",
        "objective": (
            "Deploy and test three serverless / managed services: an AWS Lambda "
            "function that writes records to DynamoDB on each invocation, the "
            "DynamoDB table itself, and an ECS Fargate task running an Nginx "
            "container - all without managing any underlying servers."
        ),
        "services": [
            "AWS Lambda",
            "Amazon DynamoDB",
            "Amazon ECS with the Fargate launch type",
            "AWS IAM (execution roles)",
            "Amazon CloudWatch Logs",
        ],
        "procedure": [
            ("Step 1: Create the DynamoDB table", [
                "Open the DynamoDB console.",
                f"Create a table named visitors-{s.suffix}.",
                "Set the partition key to id (String).",
                "Choose On-demand (PAY_PER_REQUEST) capacity mode.",
            ]),
            ("Step 2: Create IAM role and Lambda function", [
                f"Create an IAM role role-{s.suffix}-lambda with a trust policy for "
                "lambda.amazonaws.com.",
                "Attach AWSLambdaBasicExecutionRole and a custom inline policy "
                f"allowing dynamodb:PutItem on the visitors-{s.suffix} table.",
                "Open the Lambda console and create a function named "
                f"lambda-{s.suffix}-visitor-logger using Python 3.12.",
                f"Use role-{s.suffix}-lambda as the execution role.",
                f"Set environment variable TABLE_NAME = visitors-{s.suffix}.",
                "Paste in the handler.py code that builds an item and calls "
                "ddb.put_item().",
            ]),
            ("Step 3: Test the Lambda function", [
                "From the Lambda console click the Test tab.",
                "Configure a test event with body {} and save it.",
                "Click Test - the response should be a 200 with a JSON body that "
                "contains the new id and the table name it wrote to.",
                "Click Test two more times so that DynamoDB ends up with three rows.",
            ]),
            ("Step 4: Verify items in DynamoDB", [
                f"Open DynamoDB → Tables → visitors-{s.suffix}.",
                "Click Explore table items. Three items should appear, each with a "
                f"unique UUID id, a numeric timestamp, name = {s.display_name} and source = "
                f"lambda-{s.suffix}-visitor-logger.",
            ]),
            ("Step 5: Run a Fargate task", [
                f"Create an ECS cluster fargate-{s.suffix}-cluster.",
                f"Create a task definition nginx-{s.suffix}: Fargate, 0.25 vCPU, 512 MB, "
                "container image public.ecr.aws/nginx/nginx:alpine.",
                "Run a one-off task on the cluster, in the default VPC subnets, with "
                "Auto-assign public IP turned on.",
                "Wait until the task moves to Last status = RUNNING and view the "
                "CloudWatch logs to confirm the Nginx container started.",
                "Stop the task manually after capturing the screenshots.",
            ]),
        ],
        "screenshots": [
            ("lab-07/7.1-lambda-function.png",
             f"Figure 7.1 — Lambda function lambda-{s.suffix}-visitor-logger with Python 3.12 runtime."),
            ("lab-07/7.2-lambda-test-success.png",
             "Figure 7.2 — Lambda Test tab showing a successful invocation with the response body."),
            ("lab-07/7.3-dynamodb-items.png",
             f"Figure 7.3 — DynamoDB visitors-{s.suffix} table with three items written by the Lambda."),
            ("lab-07/7.4-fargate-task-running.png",
             f"Figure 7.4 — Fargate task in the RUNNING state on the fargate-{s.suffix}-cluster."),
            ("lab-07/7.5-fargate-task-logs.png",
             "Figure 7.5 — CloudWatch logs from the Fargate task showing Nginx startup messages."),
        ],
        "observations": (
            "All three serverless services worked end to end. The Lambda function, "
            "running on the AWS-managed runtime, successfully assumed its execution "
            "role, wrote items into DynamoDB on each invocation, and returned a 200 "
            "response. DynamoDB stored every record reliably with no provisioning "
            "required, since the table was created in on-demand mode. The ECS "
            "Fargate task pulled the public Nginx image from ECR Public, started "
            "the container without any EC2 host being managed by the user, and "
            "streamed its stdout into CloudWatch Logs - which is the typical "
            "developer experience for containerised serverless workloads on AWS."
        ),
    }


def lab_8(cfg):
    s = cfg.student
    return {
        "title": "Messaging Queue and Pub/Sub Simulation (SNS, SQS)",
        "objective": (
            "Configure an SNS topic with two subscribers - an SQS queue and an email "
            "address - to demonstrate the publish/subscribe (fan-out) messaging "
            "pattern, where a single published message is delivered to every "
            "confirmed subscriber."
        ),
        "services": [
            "Amazon SNS (Simple Notification Service)",
            "Amazon SQS (Simple Queue Service)",
            "Email subscription protocol",
            "SQS access policy (queue policy)",
        ],
        "procedure": [
            ("Step 1: Create the SNS topic", [
                "Open the SNS console.",
                f"Create a Standard topic named sns-{s.suffix}-notifications.",
                "Leave the rest of the settings as default.",
            ]),
            ("Step 2: Create the SQS queue", [
                "Open the SQS console.",
                f"Create a Standard queue named sqs-{s.suffix}-orders.",
                "Apply a queue access policy that allows the SNS topic to call "
                "sqs:SendMessage on this queue (using SourceArn = topic ARN).",
            ]),
            ("Step 3: Subscribe the SQS queue to the topic", [
                "From the SNS topic page, click Create subscription.",
                f"Protocol: Amazon SQS, Endpoint: ARN of sqs-{s.suffix}-orders.",
                "Click Create. The subscription becomes Confirmed automatically.",
            ]),
            ("Step 4: Subscribe an email address", [
                "Create a second subscription on the same topic.",
                "Protocol: Email, Endpoint: a real inbox address.",
                "Wait for the AWS Notification - Subscription Confirmation email.",
                "Open it and click the Confirm subscription link.",
                "The subscription's Status changes to Confirmed.",
            ]),
            ("Step 5: Publish and verify fan-out", [
                "Back on the topic, click Publish message.",
                f"Set a Subject (e.g., 'Test from {s.display_name} - Lab 8') and a Message body.",
                "Click Publish.",
                "Within a few seconds the email subscriber receives the message.",
                f"In the SQS console, open sqs-{s.suffix}-orders and click Poll for messages. "
                "The same message arrives wrapped in an SNS notification envelope.",
            ]),
        ],
        "screenshots": [
            ("lab-08/8.1-sns-topic.png",
             f"Figure 8.1 — SNS topic sns-{s.suffix}-notifications with two confirmed subscriptions."),
            ("lab-08/8.2-email-received.png",
             "Figure 8.2 — The published message delivered to the email subscriber."),
            ("lab-08/8.3-sqs-poll-result.png",
             "Figure 8.3 — The SQS queue receiving the same message wrapped in an SNS envelope."),
        ],
        "observations": (
            "The pub/sub fan-out pattern worked exactly as expected. A single Publish "
            "from the SNS topic was delivered both to the email subscriber and to "
            "the SQS queue - the SQS message arrived inside the standard SNS JSON "
            "envelope that includes Type, MessageId, TopicArn, Subject and the "
            "original Message field. This is the canonical decoupling pattern used "
            "in real AWS systems: a producer publishes once, and many independent "
            "consumers (queues, email, HTTPS endpoints, Lambda, mobile push) can "
            "process the event in their own time, without the producer needing to "
            "know about any of them."
        ),
    }


def lab_9(cfg):
    s = cfg.student
    return {
        "title": "Infrastructure as Code using CloudFormation",
        "objective": (
            "Provision AWS infrastructure declaratively using AWS CloudFormation. "
            "A YAML template describes the desired resources (a VPC and an S3 "
            "bucket); CloudFormation then handles their creation, ordering and "
            "tracking as a single managed unit called a stack."
        ),
        "services": [
            "AWS CloudFormation",
            "Amazon VPC (managed by CloudFormation)",
            "Amazon S3 (managed by CloudFormation)",
            "AWS CLI",
        ],
        "procedure": [
            ("Step 1: Author the template", [
                "Create a YAML file stack.yaml.",
                "Add the AWSTemplateFormatVersion and a Description.",
                f"Define two Parameters - OwnerName (default '{s.suffix}') and VpcCidr "
                "(default 10.50.0.0/16).",
                "Define two Resources - an AWS::EC2::VPC named CfnVpc with CIDR "
                "from the parameter, and an AWS::S3::Bucket named CfnBucket using "
                "!Sub with the AWS::AccountId pseudo-parameter for uniqueness.",
                "Define four Outputs - VpcId, VpcCidrOutput, BucketName and BucketArn.",
            ]),
            ("Step 2: Deploy the stack", [
                "Run aws cloudformation deploy --template-file stack.yaml "
                f"--stack-name stack-{s.suffix} --region {cfg.aws.region} --capabilities "
                "CAPABILITY_NAMED_IAM.",
                "Wait for the message Successfully created/updated stack.",
                "Alternatively, the same template can be uploaded through the "
                "CloudFormation console UI.",
            ]),
            ("Step 3: Inspect the stack in the console", [
                f"Open the CloudFormation console for region {cfg.aws.region}.",
                f"Locate stack-{s.suffix} in the Stacks list (Status = CREATE_COMPLETE).",
                "Open the stack and review the Events tab for the lifecycle of "
                "each resource.",
                "Open the Resources tab to see the two logical IDs (CfnVpc, CfnBucket) "
                "mapped to their actual physical IDs.",
                "Open the Outputs tab to see the values returned by the stack.",
            ]),
            ("Step 4: Delete the stack", [
                f"Run aws cloudformation delete-stack --stack-name stack-{s.suffix}.",
                "Wait until the stack disappears from the console - both the VPC "
                "and the S3 bucket are deleted by CloudFormation in the correct order.",
            ]),
        ],
        "screenshots": [
            ("lab-09/9.1-stack-list.png",
             f"Figure 9.1 — CloudFormation stack stack-{s.suffix} in CREATE_COMPLETE state."),
            ("lab-09/9.2-stack-events.png",
             "Figure 9.2 — Stack events showing CREATE_IN_PROGRESS to CREATE_COMPLETE for each resource."),
            ("lab-09/9.3-stack-resources.png",
             "Figure 9.3 — Stack resources panel: logical IDs mapped to physical resource IDs."),
            ("lab-09/9.4-stack-outputs.png",
             "Figure 9.4 — Stack outputs: VpcId, VpcCidrOutput, BucketName and BucketArn."),
        ],
        "observations": (
            "CloudFormation deployed the entire stack from a single YAML file. The "
            "Events tab confirmed that resources were created in the correct order "
            "and the stack reached CREATE_COMPLETE without manual intervention. The "
            "Outputs tab exposed values such as the VPC ID and the bucket name, "
            "which other stacks (or operators) can consume. Compared with the "
            "earlier labs that used Terraform, CloudFormation offers very similar "
            "Infrastructure as Code benefits - the template is version-controllable, "
            "repeatable across environments and self-documenting. The main "
            "difference is that CloudFormation is AWS-native and Terraform is "
            "cross-cloud, but for a single-cloud workload either is a good choice."
        ),
    }


def get_labs(cfg):
    return [lab_1(cfg), lab_2(cfg), lab_3(cfg), lab_4(cfg),
            lab_5(cfg), lab_6(cfg), lab_7(cfg), lab_8(cfg), lab_9(cfg)]


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------

def build():
    cfg = load()
    doc = Document()

    configure_default_styles(doc)
    configure_heading_styles(doc)
    configure_sections(doc, cfg)

    add_cover_page(doc, cfg)
    add_acknowledgement(doc, cfg)
    labs = get_labs(cfg)
    add_toc(doc, labs)

    for i, lab in enumerate(labs, 1):
        add_lab(doc, i, lab, is_first=(i == 1))

    add_conclusion(doc, cfg)

    _enable_update_fields_on_open(doc)
    out = output_path(cfg)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    build()
